#!/usr/bin/env python3
"""
Concord Platform Converter  (Final)
===================================

Converts a Concord Activity Player URL into curriculum-ready documents.

Outputs
-------
- PDF  (always): every page rendered and merged, with an "Answer Key" block
  injected under each multiple-choice question that has authored feedback.
- Google Docs version (optional): a structured .docx built directly from the
  activity JSON — title, activity description, page headings, text, images
  (downloaded from their URLs), multiple-choice questions with options and
  per-choice feedback, and open-response questions. Standard Letter pages that
  open cleanly and editably in Google Docs.
- Word DOCX (optional): a layout-preserving .docx made from the rendered PDF
  (online high-quality converter, with an offline fallback). Best in Microsoft
  Word when an exact visual copy is wanted.

Requirements
------------
    pip install selenium webdriver-manager requests PyPDF2 python-docx Pillow
    # optional, only for the offline Word-DOCX fallback:
    pip install pdf2docx

Chrome must be installed (used for the PDF and for fetching images that block
plain downloads). The Google Docs version is built from the JSON.
"""

import os
import re
import sys
import json
import time
import base64
import shutil
import logging
import tempfile
import threading
import html as _html

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from urllib.parse import urlparse, parse_qs, unquote, quote

import requests

from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager

try:
    from pypdf import PdfMerger
except ImportError:  # pragma: no cover
    from PyPDF2 import PdfMerger

logging.getLogger("urllib3").setLevel(logging.ERROR)


# ──────────────────────────────────────────────────────────────────────────
# Pure helpers (no GUI / no Selenium) — unit-testable on their own.
# ──────────────────────────────────────────────────────────────────────────
_TAG_RE = re.compile(r"<[^>]+>")


def html_to_text(s: str) -> str:
    """Strip HTML tags and decode entities to readable plain text."""
    if not s:
        return ""
    s = _TAG_RE.sub("", s)
    s = _html.unescape(s)
    return re.sub(r"\s+", " ", s).strip()


def _kget(d: dict, key: str, default=None):
    """Get a dict value tolerant of keys that have stray leading/trailing
    whitespace (a known artifact in some exported/observed JSON)."""
    if not isinstance(d, dict):
        return default
    if key in d:
        return d[key]
    for k in d:
        if isinstance(k, str) and k.strip() == key:
            return d[k]
    return default


def parse_mcq(node: dict):
    """
    If `node` is a ManagedInteractive multiple-choice question, return a tidy
    dict describing it; otherwise None. `authored_state` is a JSON-encoded
    *string*, so it is parsed a second time. `has_feedback` is True only when
    at least one choice carries authored choiceFeedback.
    """
    if not isinstance(node, dict) or node.get("type") != "ManagedInteractive":
        return None
    raw = node.get("authored_state")
    if not isinstance(raw, str):
        return None
    try:
        state = json.loads(raw)
    except (json.JSONDecodeError, ValueError):
        return None
    if not isinstance(state, dict) or state.get("questionType") != "multiple_choice":
        return None
    ref_id = node.get("ref_id")
    if not ref_id:
        return None

    choices = []
    for c in (_kget(state, "choices") or []):
        if isinstance(c, dict):
            choices.append({
                "content": html_to_text(_kget(c, "content", "")),
                "correct": bool(_kget(c, "correct", False)),
                "feedback": html_to_text(_kget(c, "choiceFeedback", "")),
            })
    return {
        "ref_id": ref_id,
        "prompt": html_to_text(_kget(state, "prompt", "")),
        "multiple_answers": bool(_kget(state, "multipleAnswers", False)),
        "has_feedback": any(c["feedback"] for c in choices),
        "has_correct": any(c["correct"] for c in choices),
        "choices": choices,
    }


def collect_mcqs(node, found=None) -> dict:
    """Recursively collect {ref_id: mcq_dict} for every multiple-choice
    question anywhere in the JSON tree."""
    if found is None:
        found = {}
    if isinstance(node, dict):
        mcq = parse_mcq(node)
        if mcq:
            found[mcq["ref_id"]] = mcq
        for v in node.values():
            collect_mcqs(v, found)
    elif isinstance(node, list):
        for item in node:
            collect_mcqs(item, found)
    return found


def build_answer_key_html(mcq: dict) -> str:
    """Styled HTML block injected under a feedback-bearing question."""
    esc = _html.escape

    correct = [c["content"] for c in mcq["choices"] if c["correct"]]
    if correct:
        label = "Correct answers" if (len(correct) > 1 or mcq["multiple_answers"]) else "Correct answer"
        correct_html = (
            f'<div style="margin-bottom:8px;">'
            f'<span style="font-weight:700;">{label}:</span> {esc(", ".join(correct))}</div>'
        )
    else:
        correct_html = ""

    has_feedback = any(c["feedback"] for c in mcq["choices"])

    rows = []
    if has_feedback:
        # Full per-choice breakdown with marks and feedback text.
        for c in mcq["choices"]:
            if c["correct"]:
                mark, color = "\u2713", "#2e7d32"   # check
            else:
                mark, color = "\u2717", "#c62828"   # cross
            fb = f' &mdash; {esc(c["feedback"])}' if c["feedback"] else ""
            rows.append(
                f'<li style="margin:3px 0;">'
                f'<span style="color:{color};font-weight:700;">{mark}</span> '
                f'<span style="font-weight:600;">{esc(c["content"])}</span>'
                f'<span style="color:#444;">{fb}</span></li>'
            )

    list_html = ('<ul style="margin:4px 0 0 0;padding-left:18px;list-style:none;">'
                 + "".join(rows) + '</ul>') if rows else ""

    return (
        '<div data-answer-key="1" style="'
        'border:2px solid #2e7d32;border-radius:8px;padding:10px 14px;'
        'margin:8px 0 22px 0;background:#f1f8e9;'
        'font-family:Arial,Helvetica,sans-serif;font-size:13px;color:#1b1b1b;'
        'page-break-inside:avoid;-webkit-print-color-adjust:exact;print-color-adjust:exact;">'
        '<div style="font-weight:700;color:#2e7d32;margin-bottom:6px;font-size:14px;">'
        'Answer Key</div>'
        + correct_html
        + list_html +
        '</div>'
    )


# ──────────────────────────────────────────────────────────────────────────
# JSON content extraction for the structured DOCX.
# Tree: activities[] -> pages[] -> sections[] -> embeddables[]
#   Embeddable::Xhtml         -> HTML text (may contain <img src>)
#   ManagedInteractive        -> authored_state JSON string:
#         questionType == multiple_choice  -> prompt + choices(+feedback)
#         questionType == open_response     -> prompt + defaultAnswer
#         questionType == iframe_interactive-> interactive placeholder
# ──────────────────────────────────────────────────────────────────────────
_IMG_RE = re.compile(r'<img\b[^>]*?\bsrc\s*=\s*["\']([^"\']+)["\'][^>]*>', re.I)
_TABLE_RE = re.compile(r'(?is)<table\b.*?</table>')
_IFRAME_RE = re.compile(r'(?is)<iframe\b[^>]*?\bsrc\s*=\s*["\']([^"\']+)["\'][^>]*?>\s*</iframe>')
_LINK_RE = re.compile(r'(?is)<a\b[^>]*?\bhref\s*=\s*["\']([^"\']+)["\'][^>]*?>(.*?)</a>')


def _youtube_watch_url(src: str) -> str:
    """Turn a YouTube embed URL into a normal watchable link, preserving start time."""
    src = _html.unescape(src or "")
    m = re.search(r'youtube(?:-nocookie)?\.com/embed/([\w-]+)', src)
    if not m:
        return src
    vid = m.group(1)
    start = ""
    sm = re.search(r'[?&]start=(\d+)', src)
    if sm:
        start = f"&t={sm.group(1)}s"
    return f"https://www.youtube.com/watch?v={vid}{start}"
_TR_RE = re.compile(r'(?is)<tr\b[^>]*>(.*?)</tr>')
_CELL_RE = re.compile(r'(?is)<(t[dh])\b([^>]*)>(.*?)</\1>')

# Named HTML colors used in the authored tables -> hex for Word shading.
_COLOR_HEX = {
    "lightgrey": "D9D9D9", "lightgray": "D9D9D9", "grey": "BFBFBF", "gray": "BFBFBF",
    "lightgreen": "C6EFCE", "lightyellow": "FFF2CC", "lightblue": "BDD7EE",
    "lightcyan": "DEEBF7", "white": "FFFFFF", "yellow": "FFF2CC", "green": "C6EFCE",
}


def _cell_fill(attrs: str):
    """Extract a fill color (hex) from a cell's bgcolor=/style=background attributes."""
    m = re.search(r'bgcolor\s*=\s*["\']?([#\w]+)', attrs, re.I)
    if not m:
        m = re.search(r'background(?:-color)?\s*:\s*([#\w]+)', attrs, re.I)
    if not m:
        return None
    val = m.group(1).strip().lower()
    if val.startswith("#"):
        return val[1:].upper().ljust(6, "0")[:6]
    return _COLOR_HEX.get(val)


def parse_html_table(table_html: str):
    """Parse a <table> into {'rows': [[{text,fill,header}, ...], ...]}."""
    rows = []
    for tr in _TR_RE.findall(table_html):
        cells = []
        for tag, attrs, inner in _CELL_RE.findall(tr):
            img_m = _IMG_RE.search(inner)
            img_src = ""
            if img_m:
                img_src = _html.unescape(img_m.group(1)).strip()
                if img_src.lower().startswith("data:"):
                    img_src = ""
            cells.append({
                "text": html_to_text(inner),
                "image": img_src,
                "fill": _cell_fill(attrs),
                "header": (tag.lower() == "th"),
            })
        if cells:
            rows.append(cells)
    return {"rows": rows} if rows else None


def html_to_blocks(html: str):
    """Convert authored HTML into an ordered list of content blocks:
        {'type':'text','text':...}
        {'type':'image','src':...}
        {'type':'table','rows':[[{text,fill,header}, ...], ...]}
    Tables are pulled out and rendered as real Word tables; the HTML around
    them is processed normally for text and images, in document order."""
    if not html:
        return []

    blocks = []
    pos = 0
    for tm in _TABLE_RE.finditer(html):
        # everything before this table -> text/images
        blocks.extend(_inline_html_to_blocks(html[pos:tm.start()]))
        tbl = parse_html_table(tm.group(0))
        if tbl:
            blocks.append({"type": "table", **tbl})
        pos = tm.end()
    blocks.extend(_inline_html_to_blocks(html[pos:]))
    return blocks


def _inline_html_to_blocks(html: str):
    """Handle a table-free HTML fragment: text paragraphs (with hyperlinks),
    inline images, and embedded videos (iframes)."""
    if not html:
        return []

    # Pull out <iframe> embeds first (YouTube etc.) -> video blocks, in order.
    blocks = []
    pos = 0
    for m in _IFRAME_RE.finditer(html):
        blocks.extend(_fragment_text_and_images(html[pos:m.start()]))
        url = _youtube_watch_url(m.group(1))
        if url:
            blocks.append({"type": "video", "url": url})
        pos = m.end()
    blocks.extend(_fragment_text_and_images(html[pos:]))
    return blocks


def _fragment_text_and_images(html: str):
    """Within an iframe-free fragment: emit text paragraphs (preserving links)
    and inline images, in document order."""
    if not html:
        return []
    s = html
    s = re.sub(r'(?i)<\s*br\s*/?\s*>', '\n', s)
    s = re.sub(r'(?i)</\s*(p|div|li|tr|h[1-6])\s*>', '\n', s)
    s = re.sub(r'(?i)<\s*li[^>]*>', '\n\u2022 ', s)

    blocks = []
    pos = 0
    for m in _IMG_RE.finditer(s):
        pre = s[pos:m.start()]
        blocks.extend(_paragraphs_with_links(pre))
        src = _html.unescape(m.group(1)).strip()
        if src and not src.lower().startswith("data:"):
            blocks.append({"type": "image", "src": src})
        pos = m.end()
    blocks.extend(_paragraphs_with_links(s[pos:]))
    return blocks


def _paragraphs_with_links(html_fragment: str):
    """Split an HTML fragment into paragraph blocks. Each block is either a plain
    {'type':'text','text':...} or, when it contains <a href> links,
    {'type':'richtext','segments':[{'text':..,'href':None|url}, ...]}."""
    if not html_fragment:
        return []
    # Split into lines on the newlines we inserted for block boundaries.
    out = []
    for line in html_fragment.split("\n"):
        if "<a " not in line.lower():
            txt = html_to_text(line)
            if txt:
                out.append({"type": "text", "text": txt})
            continue
        # Build ordered segments of plain text and links.
        segments = []
        p = 0
        for lm in _LINK_RE.finditer(line):
            pre = html_to_text(line[p:lm.start()])
            if pre:
                segments.append({"text": pre, "href": None})
            href = _html.unescape(lm.group(1)).strip()
            label = html_to_text(lm.group(2)) or href
            segments.append({"text": label, "href": href})
            p = lm.end()
        tail = html_to_text(line[p:])
        if tail:
            segments.append({"text": tail, "href": None})
        segments = [seg for seg in segments if seg["text"]]
        if segments:
            out.append({"type": "richtext", "segments": segments})
    return out


def _add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink (blue, underlined) to a python-docx paragraph.
    python-docx has no native hyperlink API, so we build the XML directly."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1155CC"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def parse_open_response(node: dict):
    """Return {'prompt','default'} for an open_response ManagedInteractive, else None."""
    if not isinstance(node, dict) or node.get("type") != "ManagedInteractive":
        return None
    raw = node.get("authored_state")
    if not isinstance(raw, str):
        return None
    try:
        st = json.loads(raw)
    except (json.JSONDecodeError, ValueError):
        return None
    if not isinstance(st, dict) or st.get("questionType") != "open_response":
        return None
    return {
        "prompt": html_to_text(_kget(st, "prompt", "")),
        "default": (_kget(st, "defaultAnswer", "") or "").strip(),
    }


def parse_iframe_interactive(node: dict):
    """Return a dict describing an iframe_interactive ManagedInteractive, else None.
    Two flavors exist:
      - drawing/annotation: has a 'prompt' and a 'backgroundImageUrl' (the image
        students annotate) -> treated as a question with an image
      - video: has a 'videoUrl' and a 'poster' image -> shown as the poster image
        plus a link to the video
    Returns {'prompt', 'image', 'video', 'caption'} (any may be '')."""
    if not isinstance(node, dict) or node.get("type") != "ManagedInteractive":
        return None
    raw = node.get("authored_state")
    if not isinstance(raw, str):
        return None
    try:
        st = json.loads(raw)
    except (json.JSONDecodeError, ValueError):
        return None
    if not isinstance(st, dict) or st.get("questionType") != "iframe_interactive":
        return None

    prompt = html_to_text(_kget(st, "prompt", ""))
    image = (_kget(st, "backgroundImageUrl", "") or "").strip()
    video = (_kget(st, "videoUrl", "") or "").strip()
    poster = (_kget(st, "poster", "") or "").strip()
    # Prefer the annotation background; otherwise the video poster frame.
    if not image and poster:
        image = poster
    return {
        "prompt": prompt,
        "image": image.replace(" ", ""),     # heal any stray spaces in URL
        "video": video.replace(" ", ""),
        "caption": (_kget(st, "creditLinkDisplayText", "") or "").strip(),
    }


def parse_image_interactive(node: dict):
    """Return {'image','caption'} for an 'Image Interactive' ManagedInteractive,
    else None. These have no questionType; the image lives in the 'url' field and
    `exportToMediaLibraryType == 'image'` (or library_interactive name is
    'Image Interactive'). These are the zoomable images in the activity."""
    if not isinstance(node, dict) or node.get("type") != "ManagedInteractive":
        return None
    raw = node.get("authored_state")
    if not isinstance(raw, str):
        return None
    try:
        st = json.loads(raw)
    except (json.JSONDecodeError, ValueError):
        return None
    if not isinstance(st, dict) or _kget(st, "questionType"):
        return None  # has a questionType -> handled elsewhere

    lib_name = ""
    lib = node.get("library_interactive")
    if isinstance(lib, dict):
        lib_name = ((lib.get("data") or {}).get("name") or "")

    is_image = (_kget(st, "exportToMediaLibraryType") == "image"
                or lib_name == "Image Interactive")
    if not is_image:
        return None

    url = (_kget(st, "url", "") or "").strip().replace(" ", "")
    if not url:
        return None
    caption = (_kget(st, "caption", "") or _kget(st, "creditLinkDisplayText", "") or "").strip()
    return {"image": url, "caption": caption}


def _embeddable_question_type(node: dict) -> str:
    """Best-effort questionType for a ManagedInteractive, '' otherwise."""
    if not isinstance(node, dict) or node.get("type") != "ManagedInteractive":
        return ""
    raw = node.get("authored_state")
    if not isinstance(raw, str):
        return ""
    try:
        return json.loads(raw).get("questionType", "") or ""
    except (json.JSONDecodeError, ValueError):
        return ""


def extract_activity_content(activity: dict):
    """Walk one activity and return an ordered list of pages, each:
        {'name': str, 'blocks': [ ... ]}
    where each block is one of:
        {'type':'text','text':str}
        {'type':'image','src':str,'caption':str}   (incl. zoomable Image Interactives)
        {'type':'mcq','mcq':{...}}            (from parse_mcq)
        {'type':'open_response','prompt':str,'default':str}
        {'type':'iframe','prompt':str,'image':str,'video':str,'caption':str}
        {'type':'interactive','label':str}    (other widget placeholder)
    Hidden pages and hidden embeddables are skipped. Order follows 'position'."""
    pages_out = []
    pages = sorted(activity.get("pages", []) or [],
                   key=lambda p: p.get("position", 0))
    for p in pages:
        if p.get("is_hidden"):
            continue
        page_blocks = []
        sections = sorted(p.get("sections", []) or [],
                          key=lambda s: s.get("position", 0) if isinstance(s, dict) else 0)
        for sec in sections:
            if not isinstance(sec, dict) or sec.get("is_hidden"):
                continue
            embeddables = sorted(sec.get("embeddables", []) or [],
                                 key=lambda e: e.get("position", 0) if isinstance(e, dict) else 0)
            for e in embeddables:
                if not isinstance(e, dict) or e.get("is_hidden"):
                    continue
                etype = e.get("type")
                if etype == "Embeddable::Xhtml":
                    block_name = (e.get("name") or "").strip()
                    if block_name:
                        page_blocks.append({"type": "subheading", "text": block_name})
                    page_blocks.extend(html_to_blocks(e.get("content", "")))
                elif etype == "ManagedInteractive":
                    qt = _embeddable_question_type(e)
                    if qt == "multiple_choice":
                        mcq = parse_mcq(e)
                        if mcq:
                            page_blocks.append({"type": "mcq", "mcq": mcq})
                    elif qt == "open_response":
                        opn = parse_open_response(e)
                        if opn:
                            page_blocks.append({"type": "open_response", **opn})
                    elif qt == "iframe_interactive":
                        ifr = parse_iframe_interactive(e)
                        if ifr:
                            page_blocks.append({"type": "iframe", **ifr})
                    elif qt == "":
                        img = parse_image_interactive(e)
                        if img:
                            page_blocks.append({"type": "image",
                                                "src": img["image"],
                                                "caption": img.get("caption", "")})
                        else:
                            page_blocks.append({"type": "interactive",
                                                "label": "Interactive activity"})
                elif etype == "MwInteractive":
                    page_blocks.append({"type": "interactive",
                                        "label": "Interactive activity"})
                # other embeddable types (plugins) are skipped
        pages_out.append({"name": p.get("name") or "", "blocks": page_blocks})
    return pages_out


# JavaScript snippets (triple-quoted to avoid escaping pain).
_INJECT_JS = """
var refId = arguments[0];
var html  = arguments[1];
var el = document.querySelector('[id*="' + refId + '"]');
if (!el) { return false; }
el.insertAdjacentHTML('afterend', html);
return true;
"""

_ADD_STYLE_JS = """
var s = document.createElement('style');
s.setAttribute('data-injected', 'print-fix');
s.textContent = arguments[0];
document.head.appendChild(s);
"""

_SCROLL_JS = """
var c = document.querySelector("div.app[data-cy='app']");
if (c) { c.scrollTo(0, c.scrollHeight); }
window.scrollTo(0, document.body.scrollHeight);
"""

_MEASURE_JS = """
var c = document.querySelector("div.app[data-cy='app']");
var ch = c ? c.scrollHeight : 0;
return Math.max(
  document.body ? document.body.scrollHeight : 0,
  document.documentElement ? document.documentElement.scrollHeight : 0,
  ch
);
"""

_PRINT_CSS = """
@media print {
  html, body { height: auto !important; overflow: visible !important; }
  div.app[data-cy='app'] { height: auto !important; overflow: visible !important; }
  [data-answer-key] { page-break-inside: avoid; }
}
"""



# ──────────────────────────────────────────────────────────────────────────
# Application
# ──────────────────────────────────────────────────────────────────────────
class ConcordConverter:
    PAPER_WIDTH = 12.5            # inches
    MARGINS = dict(marginTop=0.5, marginBottom=0.5, marginLeft=0.5, marginRight=0.5)
    MAX_WAIT = 30                 # seconds for page load
    PX_PER_INCH = 96.0
    MAX_PAGE_INCHES = 200.0

    def __init__(self, root: tk.Tk, ui_scale: float = 1.0):
        self.root = root
        self.ui_scale = ui_scale if ui_scale and ui_scale > 0 else 1.0
        self.root.title("Concord Activity Converter")
        w = int(750 * self.ui_scale)
        h = int(720 * self.ui_scale)
        self.root.geometry(f"{w}x{h}")
        self.root.minsize(w, h)
        self.root.resizable(False, False)
        self._busy = False
        self._build_gui()

    # ── GUI ────────────────────────────────────────────────────────────────
    def _build_gui(self):
        # Color scheme
        self.BG_PRIMARY = "#F8F9FA"
        self.BG_CARD = "#FFFFFF"
        self.COLOR_PRIMARY = "#1976D2"
        self.COLOR_SUCCESS = "#4CAF50"
        self.COLOR_TEXT = "#212121"
        self.COLOR_SUBTEXT = "#757575"
        self.COLOR_BORDER = "#E0E0E0"

        self.root.configure(bg=self.BG_PRIMARY)

        # Themed widgets (ttk) — 'clam' renders crisply and lets us recolor.
        style = ttk.Style()
        try:
            style.theme_use("clam")
        except Exception:
            pass
        style.configure("Concord.Horizontal.TProgressbar",
                        troughcolor="#E8EAED", bordercolor="#E8EAED",
                        background=self.COLOR_SUCCESS, lightcolor=self.COLOR_SUCCESS,
                        darkcolor=self.COLOR_SUCCESS, thickness=int(18 * self.ui_scale))
        style.configure("Concord.TEntry", fieldbackground="white",
                        bordercolor=self.COLOR_BORDER, relief="flat",
                        padding=int(6 * self.ui_scale))

        self.root.configure(bg=self.BG_PRIMARY)

        # ── HEADER ──────────────────────────────────────────────────────────
        header = tk.Frame(self.root, bg=self.COLOR_PRIMARY, height=90)
        header.pack(fill=tk.X, side=tk.TOP)
        header.pack_propagate(False)

        title = tk.Label(
            header, text="Concord Activity Converter",
            font=("Segoe UI", 22, "bold"), fg="white", bg=self.COLOR_PRIMARY
        )
        title.pack(pady=12, padx=20, anchor=tk.W)

        subtitle = tk.Label(
            header,
            text="Activities to PDF or Google Docs",
            font=("Segoe UI", 10), fg="#E3F2FD", bg=self.COLOR_PRIMARY
        )
        subtitle.pack(padx=20, anchor=tk.W)

        # ── MAIN CONTENT ────────────────────────────────────────────────────
        main = tk.Frame(self.root, bg=self.BG_PRIMARY)
        main.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # URL Input Card
        card_url = self._create_card(main, "Concord Activity URL")
        self.url_entry = ttk.Entry(card_url, width=70, font=("Segoe UI", 10))
        self.url_entry.pack(fill=tk.X, padx=15, pady=(0, 15))

        # Output Folder Card
        card_out = self._create_card(main, "Output Folder")
        folder_frame = tk.Frame(card_out, bg=self.BG_CARD)
        folder_frame.pack(fill=tk.X, padx=15, pady=(0, 15))

        self.folder_path = tk.StringVar()
        folder_entry = ttk.Entry(folder_frame, textvariable=self.folder_path, width=55)
        folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))

        browse_btn = tk.Button(
            folder_frame, text="Browse",
            font=("Segoe UI", 10, "bold"),
            bg=self.COLOR_PRIMARY, fg="white",
            relief=tk.FLAT, cursor="hand2",
            padx=15, pady=6,
            command=self._browse_folder
        )
        browse_btn.pack(side=tk.LEFT)

        # Options Card — three independent output formats, shown as clear
        # selectable cards that highlight when chosen.
        card_opt = self._create_card(main, "Output options", bottom_pad=6)

        intro = tk.Label(
            card_opt, text="Click to select what to create (selected = blue):",
            font=("Segoe UI", 9), fg=self.COLOR_SUBTEXT, bg=self.BG_CARD
        )
        intro.pack(anchor=tk.W, padx=15, pady=(0, 8))

        self.pdf_var = tk.BooleanVar(value=True)
        self.gdocs_var = tk.BooleanVar(value=True)
        self.worddoc_var = tk.BooleanVar(value=False)

        self._make_option_toggle(card_opt, self.pdf_var, "PDF")

        # PDF sub-options: with vs without answers/feedback (mutually exclusive)
        self.pdf_answers_var = tk.StringVar(value="with")
        sub = tk.Frame(card_opt, bg=self.BG_CARD)
        sub.pack(anchor=tk.W, fill=tk.X, padx=(46, 15), pady=(2, 6))
        for label, val in (("With answers and choice feedback", "with"),
                           ("Without answers or feedback", "without")):
            tk.Radiobutton(sub, text=label, variable=self.pdf_answers_var, value=val,
                           font=("Segoe UI", 10), bg=self.BG_CARD, fg=self.COLOR_TEXT,
                           activebackground=self.BG_CARD, selectcolor="white",
                           highlightthickness=0).pack(anchor=tk.W)

        self._make_option_toggle(card_opt, self.gdocs_var, "Google Docs version")
        self._make_option_toggle(card_opt, self.worddoc_var,
                                 "Word DOCX   (preserves the original page layout)")

        # Status / Progress Card — clean progress bar only (no log strings)
        card_status = self._create_card(main, "Progress")

        self.status_label = tk.Label(
            card_status, text="Ready",
            font=("Segoe UI", 10), fg=self.COLOR_SUBTEXT, bg=self.BG_CARD
        )
        self.status_label.pack(anchor=tk.W, padx=15, pady=(0, 10))

        self.progress_bar = ttk.Progressbar(
            card_status, length=300, mode='determinate', value=0,
            style="Concord.Horizontal.TProgressbar"
        )
        self.progress_bar.pack(fill=tk.X, padx=15, pady=(0, 16))

        # Convert & Close Buttons (Footer)
        button_frame = tk.Frame(self.root, bg=self.BG_PRIMARY)
        button_frame.pack(side=tk.BOTTOM, pady=15)

        self.convert_btn = tk.Button(
            button_frame, text="Convert",
            font=("Segoe UI", 12, "bold"),
            bg=self.COLOR_PRIMARY, fg="white",
            relief=tk.FLAT, cursor="hand2",
            padx=40, pady=12,
            command=self.start_conversion,
            activebackground="#1565C0"
        )
        self.convert_btn.pack(side=tk.LEFT, padx=10)

        close_btn = tk.Button(
            button_frame, text="Close",
            font=("Segoe UI", 12, "bold"),
            bg="#757575", fg="white",
            relief=tk.FLAT, cursor="hand2",
            padx=40, pady=12,
            command=self.root.destroy,
            activebackground="#616161"
        )
        close_btn.pack(side=tk.LEFT, padx=10)

    def _create_card(self, parent, title, bottom_pad=15):
        """Create a styled card (frame with title)."""
        card = tk.Frame(
            parent, bg=self.BG_CARD, relief=tk.FLAT,
            highlightthickness=1, highlightbackground=self.COLOR_BORDER
        )
        card.pack(fill=tk.X, pady=(0, bottom_pad))

        label = tk.Label(
            card, text=title,
            font=("Segoe UI", 11, "bold"),
            fg=self.COLOR_TEXT, bg=self.BG_CARD
        )
        label.pack(anchor=tk.W, padx=15, pady=(12, 10))

        return card

    def _make_option_toggle(self, parent, var, text):
        """A clickable option row that clearly highlights when selected:
        blue background + checkmark when on, light grey when off."""
        ON_BG, ON_FG = self.COLOR_PRIMARY, "white"
        OFF_BG, OFF_FG = "#EEF1F4", self.COLOR_TEXT

        row = tk.Frame(parent, bg=self.BG_CARD)
        row.pack(fill=tk.X, padx=15, pady=3)

        box = tk.Frame(row, bg=OFF_BG, highlightthickness=1,
                       highlightbackground="#CBD2D9", cursor="hand2")
        box.pack(fill=tk.X)

        mark = tk.Label(box, text="\u2713", font=("Segoe UI", 12, "bold"),
                        width=2, bg=OFF_BG, fg=OFF_BG)  # hidden when off
        mark.pack(side=tk.LEFT, padx=(10, 2), pady=8)
        lbl = tk.Label(box, text=text, font=("Segoe UI", 11),
                       bg=OFF_BG, fg=OFF_FG, anchor="w", cursor="hand2")
        lbl.pack(side=tk.LEFT, fill=tk.X, expand=True, pady=8)

        def refresh():
            on = var.get()
            bg = ON_BG if on else OFF_BG
            fg = ON_FG if on else OFF_FG
            for w in (box, mark, lbl):
                w.configure(bg=bg)
            lbl.configure(fg=fg)
            mark.configure(fg=("white" if on else bg))  # check visible only when on
            box.configure(highlightbackground=(ON_BG if on else "#CBD2D9"))

        def toggle(_evt=None):
            var.set(not var.get())
            refresh()

        for w in (box, mark, lbl):
            w.bind("<Button-1>", toggle)
        refresh()
        return row

    def _browse_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.folder_path.set(folder)

    def log(self, msg: str):
        """Internal progress messages go to the console only, never to the GUI,
        so the progress area stays clean (no URLs or step-by-step strings)."""
        try:
            print(str(msg))
        except Exception:
            pass

    def set_status(self, text: str, color=None):
        """Set the single clean status word shown above the progress bar."""
        self.root.after(0, lambda: self.status_label.config(
            text=text, fg=color or self.COLOR_SUBTEXT))

    # ── small utilities ─────────────────────────────────────────────────────
    @staticmethod
    def sanitize_filename(name: str) -> str:
        name = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", name or "").strip()
        return name or "Activity"

    @staticmethod
    def ensure_dirs(*paths: str):
        for p in paths:
            if p:
                os.makedirs(p, exist_ok=True)

    @staticmethod
    def get_unique_filepath(path: str) -> str:
        base, ext = os.path.splitext(path)
        counter, candidate = 1, path
        while os.path.exists(candidate):
            candidate = f"{base}({counter}){ext}"
            counter += 1
        return candidate

    @staticmethod
    def _find_chrome_binary():
        """Locate the installed Chrome/Chromium executable across platforms.
        Returns a path string, or None to let Selenium auto-detect."""
        candidates = []
        if sys.platform == "darwin":  # macOS
            candidates = [
                "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
                "/Applications/Google Chrome Beta.app/Contents/MacOS/Google Chrome Beta",
                "/Applications/Google Chrome Dev.app/Contents/MacOS/Google Chrome Dev",
                "/Applications/Google Chrome Canary.app/Contents/MacOS/Google Chrome Canary",
                "/Applications/Chromium.app/Contents/MacOS/Chromium",
                os.path.expanduser(
                    "~/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"),
                os.path.expanduser(
                    "~/Applications/Google Chrome Dev.app/Contents/MacOS/Google Chrome Dev"),
            ]
        elif sys.platform.startswith("win"):  # Windows
            pf = os.environ.get("PROGRAMFILES", r"C:\Program Files")
            pf86 = os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)")
            local = os.environ.get("LOCALAPPDATA", "")
            candidates = [
                os.path.join(pf, "Google", "Chrome", "Application", "chrome.exe"),
                os.path.join(pf86, "Google", "Chrome", "Application", "chrome.exe"),
                os.path.join(local, "Google", "Chrome", "Application", "chrome.exe"),
            ]
        else:  # Linux
            candidates = [
                "/usr/bin/google-chrome",
                "/usr/bin/google-chrome-stable",
                "/usr/bin/chromium-browser",
                "/usr/bin/chromium",
            ]
        for path in candidates:
            if path and os.path.exists(path):
                return path
        return None

    def get_driver(self) -> webdriver.Chrome:
        opts = Options()
        opts.add_argument("--headless=new")
        opts.add_argument("--disable-gpu")
        opts.add_argument("--no-sandbox")
        opts.add_argument("--disable-dev-shm-usage")
        opts.add_argument("--window-size=1920,1080")
        opts.add_argument("--disable-extensions")
        opts.add_argument("--disable-software-rasterizer")
        opts.add_argument("--mute-audio")
        opts.add_experimental_option("excludeSwitches", ["enable-logging"])

        chrome_path = self._find_chrome_binary()
        if chrome_path:
            opts.binary_location = chrome_path
            self.log(f"[info] using Chrome at: {chrome_path}")

        try:
            return webdriver.Chrome(
                service=Service(ChromeDriverManager().install()), options=opts)
        except Exception as e:
            if not chrome_path:
                raise RuntimeError(
                    "Google Chrome was not found. Please install Google Chrome from "
                    "https://www.google.com/chrome and try again."
                ) from e
            raise

    # ── URL / metadata ───────────────────────────────────────────────────────
    def extract_activity_info(self, url: str) -> dict:
        parsed = urlparse(url)
        qs = parse_qs(parsed.query, keep_blank_values=True)

        if "activity" in qs:
            act_url = unquote(qs["activity"][0])
            m = re.search(r"activities/(\d+)\.json", act_url) or re.search(r"/(\d+)\.json", act_url)
            return dict(
                format="activity", activity_url=act_url,
                activity_id=(m.group(1) if m else None),
                run_key=qs.get("runKey", [None])[0], page=qs.get("page", [None])[0],
            )

        if "sequence" in qs:
            seq_url = unquote(qs["sequence"][0])
            m = re.search(r"sequences/(\d+)\.json", seq_url) or re.search(r"/(\d+)\.json", seq_url)
            return dict(
                format="sequence", sequence_url=seq_url,
                sequence_id=(m.group(1) if m else None),
                activity_id=qs.get("sequenceActivity", [""])[0],
                preview="preview" in qs, page=qs.get("page", [None])[0],
            )

        raise ValueError("URL must contain an 'activity' or 'sequence' parameter.")

    def build_page_url(self, info: dict, page_id=None) -> str:
        base = "https://activity-player.concord.org/?"
        if info["format"] == "activity":
            params = {
                "activity": quote(info["activity_url"], safe=""),
                "page": page_id, "runKey": info.get("run_key"),
            }
            params = {k: v for k, v in params.items() if v is not None}
            return base + "&".join(f"{k}={v}" for k, v in params.items())

        params = {
            "page": page_id, "preview": "",
            "sequence": quote(info["sequence_url"], safe=""),
            "sequenceActivity": info["activity_id"],
        }
        params = {k: v for k, v in params.items() if v is not None}
        return base + "&".join(k if v == "" else f"{k}={v}" for k, v in params.items())

    def fetch_source_json(self, info: dict) -> dict:
        url = info["sequence_url"] if info["format"] == "sequence" else info["activity_url"]
        if not url:
            raise ValueError("Could not determine the source JSON URL from the input.")
        resp = requests.get(url, timeout=30, headers={"User-Agent": "ConcordConverter/3.0"})
        resp.raise_for_status()
        return resp.json()

    def derive_metadata(self, info: dict, data: dict):
        if info["format"] == "sequence":
            target = (info.get("activity_id") or "").removeprefix("activity_")
            activities = data.get("activities", []) or []
            act = next((a for a in activities if str(a.get("id")) == target), None)
            if act is None:
                if len(activities) == 1:
                    act = activities[0]
                else:
                    raise ValueError(f"Activity '{info.get('activity_id')}' not found in sequence.")
            return (act.get("name") or f"Activity_{info.get('activity_id')}"), (act.get("pages") or [])
        return (data.get("name") or "Activity"), (data.get("pages") or [])

    # ── rendering ─────────────────────────────────────────────────────────────
    def inject_answer_keys(self, driver, mcq_map: dict) -> int:
        injected = 0
        for ref_id, mcq in mcq_map.items():
            # Show a box if the question has feedback OR a marked correct answer.
            # Skip only poll-style questions (no feedback and nothing marked correct).
            if not mcq.get("has_feedback") and not mcq.get("has_correct"):
                continue
            try:
                if driver.execute_script(_INJECT_JS, ref_id, build_answer_key_html(mcq)):
                    injected += 1
            except Exception as e:
                self.log(f"     [warn] could not inject answer key for {ref_id}: {e}")
        return injected

    def save_page_pdf(self, driver, url: str, out_path: str,
                      mcq_map: dict = None, include_feedback: bool = True) -> bool:
        """Render a single page to PDF (with answer keys injected). Returns success."""
        try:
            self.log(f"  -> rendering: {url}")
            driver.get(url)
            WebDriverWait(driver, self.MAX_WAIT).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, "div.app[data-cy='app']"))
            )
            time.sleep(2)

            for _ in range(15):
                driver.execute_script(_SCROLL_JS)
                time.sleep(0.25)
            driver.execute_script("window.scrollTo(0, 0);")
            time.sleep(0.3)

            driver.execute_script(_ADD_STYLE_JS, _PRINT_CSS)

            if include_feedback and mcq_map:
                n = self.inject_answer_keys(driver, mcq_map)
                if n:
                    self.log(f"     injected {n} answer-key block(s)")
                time.sleep(0.4)

            height_px = driver.execute_script(_MEASURE_JS) or 1000
            height_in = (height_px / self.PX_PER_INCH
                         + self.MARGINS["marginTop"] + self.MARGINS["marginBottom"] + 0.5)
            height_in = min(max(height_in, 4.0), self.MAX_PAGE_INCHES)

            header = (f"<div style='font-size:9px;width:100%;text-align:center;color:#666;'>"
                      f"{_html.escape(os.path.splitext(os.path.basename(out_path))[0])}</div>")
            pdf = driver.execute_cdp_cmd("Page.printToPDF", {
                "printBackground": True,
                "preferCSSPageSize": False,
                "paperWidth": self.PAPER_WIDTH,
                "paperHeight": height_in,
                **self.MARGINS,
                "displayHeaderFooter": True,
                "headerTemplate": header,
                "footerTemplate": (
                    "<div style='font-size:10px;text-align:center;width:100%;'>"
                    "<span class='pageNumber'></span> / <span class='totalPages'></span></div>"
                ),
            })
            with open(out_path, "wb") as f:
                f.write(base64.b64decode(pdf["data"]))
            return True

        except Exception as e:
            self.log(f"  [error] page render failed: {e}")
            return False

    def render_and_merge(self, driver, info, pages, safe_title, out_dir,
                         temp_dir, mcq_map, include_feedback) -> str:
        """Render all pages to PDF and merge into one file. Returns the PDF path."""
        temp_files = []

        home_pdf = os.path.join(temp_dir, f"{safe_title}_home.pdf")
        if self.save_page_pdf(driver, self.build_page_url(info), home_pdf,
                              mcq_map, include_feedback):
            temp_files.append(home_pdf)

        for idx, pg in enumerate(pages, start=1):
            pid = str(pg.get("id", idx))
            page_param = pid if pid.startswith("page_") else f"page_{pid}"
            pdf_path = os.path.join(temp_dir, f"{safe_title}_p{idx:03d}.pdf")
            url = self.build_page_url(info, page_param)
            if self.save_page_pdf(driver, url, pdf_path, mcq_map, include_feedback):
                temp_files.append(pdf_path)

        if not temp_files:
            raise RuntimeError("No pages were rendered; aborting merge.")

        final_pdf = self.get_unique_filepath(os.path.join(out_dir, f"{safe_title}.pdf"))
        merger = PdfMerger()
        try:
            for fpath in temp_files:
                merger.append(fpath)
            with open(final_pdf, "wb") as out_f:
                merger.write(out_f)
        finally:
            merger.close()
        return final_pdf

    # ── DOCX (structured, built directly from the activity JSON) ────────────────
    def _image_for_docx(self, src: str, driver=None):
        """Return (BytesIO, width_px) for an image URL, or None.
        Tries a normal HTTP download first (with browser-like headers); if that
        fails and a Selenium driver is available, fetches the image through
        Chrome itself (which can access URLs that block plain HTTP clients).
        Skips SVGs (python-docx can't embed them)."""
        import io
        if not src or src.lower().startswith("data:"):
            return None
        if src.lower().split("?")[0].endswith(".svg"):
            return None

        raw = self._download_image_http(src)
        if raw is None and driver is not None:
            self.log(f"     fetching image via browser: {src[:60]}")
            raw = self._download_image_browser(driver, src)
        if not raw or len(raw) < 200:
            if raw is None:
                self.log(f"     [warn] image could not be downloaded: {src[:60]}")
            return None

        return self._normalize_image_bytes(raw)

    def _download_image_http(self, src: str):
        """Plain HTTP download with browser-like headers. Returns bytes or None."""
        try:
            headers = {
                "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                               "AppleWebKit/537.36 (KHTML, like Gecko) "
                               "Chrome/124.0.0.0 Safari/537.36"),
                "Accept": "image/avif,image/webp,image/png,image/*,*/*;q=0.8",
                "Referer": "https://activity-player.concord.org/",
            }
            resp = requests.get(src, timeout=25, headers=headers)
            resp.raise_for_status()
            if "svg" in resp.headers.get("Content-Type", "").lower():
                return None
            return resp.content
        except Exception:
            return None

    def _download_image_browser(self, driver, src: str):
        """Fetch an image using the live Chrome session. Navigating directly to
        the image makes the page same-origin with the image, so we can read it
        off a canvas as base64 PNG without CORS tainting. Returns bytes or None."""
        try:
            driver.get(src)
            time.sleep(0.8)
            data_url = driver.execute_async_script(
                """
                var cb = arguments[arguments.length - 1];
                try {
                  var img = document.images && document.images[0];
                  if (!img) { cb(null); return; }
                  var draw = function () {
                    try {
                      var c = document.createElement('canvas');
                      c.width = img.naturalWidth; c.height = img.naturalHeight;
                      c.getContext('2d').drawImage(img, 0, 0);
                      cb(c.toDataURL('image/png'));
                    } catch (e) { cb(null); }
                  };
                  if (img.complete && img.naturalWidth) { draw(); }
                  else { img.onload = draw; img.onerror = function(){ cb(null); }; }
                } catch (e) { cb(null); }
                """
            )
            if not data_url or not data_url.startswith("data:image"):
                return None
            b64 = data_url.split(",", 1)[1]
            return base64.b64decode(b64)
        except Exception:
            return None

    def _normalize_image_bytes(self, raw: bytes):
        """Re-encode image bytes to a python-docx-friendly format.
        Returns (BytesIO, width_px)."""
        import io
        try:
            from PIL import Image
        except ImportError:
            return io.BytesIO(raw), 0
        try:
            im = Image.open(io.BytesIO(raw))
            width_px = im.size[0]
            if im.mode not in ("RGB", "RGBA"):
                im = im.convert("RGB")
            out = io.BytesIO()
            fmt = "PNG" if im.mode == "RGBA" else "JPEG"
            im.save(out, format=fmt, quality=90)
            out.seek(0)
            return out, width_px
        except Exception:
            return io.BytesIO(raw), 0

    def build_docx_from_json(self, data, info, title, out_dir, driver=None):
        """Build a structured, Google Docs-compatible DOCX directly from the
        activity JSON: page headings, text, images, MCQs (with options +
        choice feedback), and open-response questions, all in document order.
        If a Selenium `driver` is supplied, it is used as a fallback to fetch
        images that block plain HTTP downloads."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise RuntimeError(
                "The 'python-docx' library is required for DOCX output.\n\n"
                "Please install it by running:\n"
                "    pip install python-docx\n\n"
                "Then run the converter again."
            )

        self.log("[info] building structured DOCX from JSON...")

        # Resolve the target activity (sequence) or the activity itself.
        if info["format"] == "sequence":
            target = (info.get("activity_id") or "").removeprefix("activity_")
            activities = data.get("activities", []) or []
            activity = next((a for a in activities if str(a.get("id")) == target), None)
            if activity is None and len(activities) == 1:
                activity = activities[0]
            if activity is None:
                raise RuntimeError("Could not locate the activity inside the sequence JSON.")
        else:
            activity = data

        pages = extract_activity_content(activity)

        docx_path = self.get_unique_filepath(
            os.path.join(out_dir, self.sanitize_filename(title) + ".docx")
        )

        doc = Document()
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, m, Inches(1))

        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(11)

        content_w_in = 6.5
        q_num = [0]  # mutable counter for question numbering

        # Title
        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def shade(cell, fill):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

        def cell_borders(cell, color):
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                e = OxmlElement(f"w:{edge}")
                e.set(qn("w:val"), "single")
                e.set(qn("w:sz"), "8")
                e.set(qn("w:color"), color)
                borders.append(e)
            tcPr.append(borders)

        def add_mcq(mcq):
            q_num[0] += 1
            # Question prompt
            p = doc.add_paragraph()
            r = p.add_run(f"Question {q_num[0]}: ")
            r.bold = True
            p.add_run(mcq.get("prompt") or "")

            multi = mcq.get("multiple_answers")
            for c in mcq["choices"]:
                cp = doc.add_paragraph(style="List Bullet")
                box = "\u2611 " if c["correct"] else "\u2610 "  # ballot box (checked/empty)
                br = cp.add_run(box)
                br.bold = True
                if c["correct"]:
                    br.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                cr = cp.add_run(c["content"])
                if c["correct"]:
                    cr.bold = True

            # Answer + feedback box (only if there is feedback OR a correct answer)
            correct = [c["content"] for c in mcq["choices"] if c["correct"]]
            has_fb = any(c["feedback"] for c in mcq["choices"])
            if correct or has_fb:
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                tbl.autofit = False
                cell = tbl.cell(0, 0)
                cell.width = Inches(content_w_in)
                shade(cell, "F1F8E9")
                cell_borders(cell, "2E7D32")
                hp = cell.paragraphs[0]
                hr = hp.add_run("Answer Key")
                hr.bold = True
                hr.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
                if correct:
                    label = ("Correct answers" if (len(correct) > 1 or multi)
                             else "Correct answer")
                    ap = cell.add_paragraph()
                    al = ap.add_run(f"{label}: ")
                    al.bold = True
                    ap.add_run(", ".join(correct))
                if has_fb:
                    for c in mcq["choices"]:
                        if not c["feedback"]:
                            continue
                        fp = cell.add_paragraph()
                        mark = "\u2713 " if c["correct"] else "\u2717 "
                        mr = fp.add_run(mark)
                        mr.bold = True
                        mr.font.color.rgb = (RGBColor(0x2E, 0x7D, 0x32) if c["correct"]
                                             else RGBColor(0xC6, 0x28, 0x28))
                        nr = fp.add_run(f"{c['content']}: ")
                        nr.bold = True
                        fp.add_run(c["feedback"])
            doc.add_paragraph()

        def add_open(blk):
            q_num[0] += 1
            p = doc.add_paragraph()
            r = p.add_run(f"Question {q_num[0]}: ")
            r.bold = True
            p.add_run(blk.get("prompt") or "")
            tbl = doc.add_table(rows=1, cols=1)
            tbl.autofit = False
            cell = tbl.cell(0, 0)
            cell.width = Inches(content_w_in)
            cell_borders(cell, "BBBBBB")
            default = (blk.get("default") or "").strip()
            if default:
                for line in default.split("\n"):
                    cp = cell.add_paragraph()
                    cp.add_run(line if line.strip() else "\u00a0")
            else:
                # blank answer space
                for _ in range(3):
                    cell.add_paragraph("\u00a0")
            doc.add_paragraph()

        def add_image(src, caption=""):
            """Download and embed an image, centered, scaled to the content width.
            Optionally adds a small italic caption beneath it."""
            got = self._image_for_docx(src, driver=driver)
            if not got:
                return False
            stream, w_px = got
            try:
                width_in = content_w_in
                if w_px:
                    width_in = min(content_w_in, w_px / 96.0)
                doc.add_picture(stream, width=Inches(width_in))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cp.add_run(caption)
                    cr.italic = True
                    cr.font.size = Pt(9)
                return True
            except Exception as e:
                self.log(f"     [warn] could not embed image: {e}")
                return False

        def _embed_image_in_cell(cell, src):
            """Download src and place it inside a table cell, scaled to fit."""
            got = self._image_for_docx(src, driver=driver)
            if not got:
                return False
            stream, w_px = got
            try:
                # cap avatar/image width so it fits a cell; ~1.1in works well
                width_in = 1.1
                if w_px:
                    width_in = min(width_in, max(0.4, w_px / 96.0))
                run = cell.paragraphs[0].add_run()
                run.add_picture(stream, width=Inches(width_in))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                return True
            except Exception as e:
                self.log(f"     [warn] could not embed cell image: {e}")
                return False

        def add_table(rows):
            """Render a parsed HTML table as a real Word table with borders,
            header shading, per-cell background colors, and in-cell images
            (e.g. conversation avatars)."""
            if not rows:
                return
            ncols = max(len(r) for r in rows)
            # A table is treated as having a header row only if its first row has
            # no images (image tables like dialogues have no header).
            has_any_image = any(c.get("image") for r in rows for c in r)
            tbl = doc.add_table(rows=0, cols=ncols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                cells = tbl.add_row().cells
                for c_idx in range(ncols):
                    cell = cells[c_idx]
                    data = row[c_idx] if c_idx < len(row) else {}
                    fill = data.get("fill")
                    if fill:
                        shade(cell, fill)
                    img_src = data.get("image")
                    if img_src and _embed_image_in_cell(cell, img_src):
                        # if there's also text, add it below the image
                        if data.get("text"):
                            tp = cell.add_paragraph()
                            tp.add_run(data["text"]).font.size = Pt(9)
                    else:
                        p = cell.paragraphs[0]
                        run = p.add_run(data.get("text", ""))
                        if data.get("header") or (r_idx == 0 and not has_any_image):
                            run.bold = True
                        run.font.size = Pt(9)
            doc.add_paragraph()

        def add_iframe(blk):
            """iframe_interactive: a prompt (numbered question) plus its image
            (annotation background or video poster), and a video link if present."""
            prompt = blk.get("prompt") or ""
            image = blk.get("image") or ""
            video = blk.get("video") or ""

            if prompt:
                q_num[0] += 1
                p = doc.add_paragraph()
                r = p.add_run(f"Question {q_num[0]}: ")
                r.bold = True
                p.add_run(prompt)

            embedded = False
            if image:
                embedded = add_image(image)
            if not embedded and not prompt and not video:
                # nothing usable; leave a small placeholder
                ip = doc.add_paragraph()
                ir = ip.add_run("[ Interactive activity ]")
                ir.italic = True
                ir.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                return

            if video:
                vp = doc.add_paragraph()
                vr = vp.add_run("Video: ")
                vr.bold = True
                vp.add_run(video)
                cap = blk.get("caption") or ""
                if cap:
                    cp = doc.add_paragraph()
                    cr = cp.add_run(cap)
                    cr.italic = True
                    cr.font.size = Pt(9)
            doc.add_paragraph()

        # Activity description + estimated time, shown under the title.
        description_html = activity.get("description") or ""
        if description_html:
            for blk in html_to_blocks(description_html):
                if blk["type"] == "text":
                    p = doc.add_paragraph(blk["text"])
                elif blk["type"] == "image":
                    add_image(blk["src"])
        ttc = activity.get("time_to_complete")
        if ttc:
            tp = doc.add_paragraph()
            tr = tp.add_run(f"Estimated time to complete: {ttc} minutes")
            tr.bold = True
            tr.italic = True
        doc.add_paragraph()

        # Build the document page by page
        for pg in pages:
            name = (pg.get("name") or "").strip()
            if name:
                doc.add_heading(name, level=1)
            for blk in pg["blocks"]:
                t = blk["type"]
                if t == "text":
                    doc.add_paragraph(blk["text"])
                elif t == "richtext":
                    rp = doc.add_paragraph()
                    for seg in blk["segments"]:
                        if seg.get("href"):
                            _add_hyperlink(rp, seg["href"], seg["text"])
                        else:
                            rp.add_run(seg["text"])
                elif t == "video":
                    vp = doc.add_paragraph()
                    vr = vp.add_run("Video: ")
                    vr.bold = True
                    _add_hyperlink(vp, blk["url"], blk["url"])
                elif t == "subheading":
                    sp = doc.add_paragraph()
                    sr = sp.add_run(blk["text"])
                    sr.bold = True
                    sr.font.size = Pt(13)
                elif t == "image":
                    add_image(blk["src"], blk.get("caption", ""))
                elif t == "table":
                    add_table(blk["rows"])
                elif t == "mcq":
                    add_mcq(blk["mcq"])
                elif t == "open_response":
                    add_open(blk)
                elif t == "iframe":
                    add_iframe(blk)
                elif t == "interactive":
                    ip = doc.add_paragraph()
                    ir = ip.add_run(f"[ {blk.get('label', 'Interactive activity')} ]")
                    ir.italic = True
                    ir.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.save(docx_path)
        return docx_path, "structured (from JSON)"

    # ── DOCX (high-fidelity, layout-preserving from the PDF) ─────────────────────
    def build_docx_highfidelity(self, pdf_path: str, title: str, out_dir: str):
        """Produce a layout-preserving DOCX from the rendered PDF. Tries the
        online high-quality converter first, then falls back to a fully offline
        converter. Returns (docx_path, engine)."""
        docx_path = self.get_unique_filepath(
            os.path.join(out_dir, self.sanitize_filename(title) + " (Word layout).docx")
        )
        try:
            self.log("[info] high-fidelity DOCX (online converter)...")
            self._convert_pdf_online(pdf_path, docx_path)
            return docx_path, "online"
        except Exception as e:
            self.log(f"[warn] online converter failed: {e}; using offline converter")
            self._convert_pdf_offline(pdf_path, docx_path)
            return docx_path, "offline"

    def _convert_pdf_online(self, pdf_path: str, docx_path: str):
        """Convert a PDF to DOCX using a free online service via Selenium."""
        driver = self.get_driver()
        try:
            driver.get("https://www.ilovepdf.com/pdf_to_word")
            time.sleep(2)
            try:
                btn = WebDriverWait(driver, 5).until(
                    EC.element_to_be_clickable((By.XPATH, "//button[contains(., 'Accept')]"))
                )
                btn.click()
                time.sleep(1)
            except Exception:
                pass
            file_input = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.XPATH, "//input[@type='file']"))
            )
            file_input.send_keys(os.path.abspath(pdf_path))
            time.sleep(2)
            try:
                WebDriverWait(driver, 30).until(
                    EC.invisibility_of_element_located(
                        (By.XPATH, "//*[contains(text(), 'processing')]"))
                )
            except Exception:
                pass
            convert_btn = WebDriverWait(driver, 20).until(
                EC.element_to_be_clickable(
                    (By.XPATH, "//button[contains(., 'Convert to WORD') or contains(., 'CONVERT')]"))
            )
            convert_btn.click()
            link = WebDriverWait(driver, 120).until(
                EC.presence_of_element_located(
                    (By.XPATH, "//a[contains(., 'Download') or contains(., 'DOWNLOAD')]"))
            )
            url = link.get_attribute("href")
            if not url:
                raise RuntimeError("no download link")
            resp = requests.get(url, timeout=60)
            resp.raise_for_status()
            with open(docx_path, "wb") as f:
                f.write(resp.content)
        finally:
            try:
                driver.quit()
            except Exception:
                pass

    def _convert_pdf_offline(self, pdf_path: str, docx_path: str):
        """Offline PDF→DOCX fallback using pdf2docx (no internet needed)."""
        try:
            from pdf2docx import Converter
        except ImportError:
            raise RuntimeError(
                "The offline DOCX converter needs 'pdf2docx'.\n\n"
                "Install it with:\n    pip install pdf2docx\n"
            )
        cv = Converter(pdf_path)
        try:
            cv.convert(docx_path)
        finally:
            cv.close()

    # ── orchestration ─────────────────────────────────────────────────────────
    def start_conversion(self):
        if self._busy:
            return
        url = self.url_entry.get().strip()
        out_dir = self.folder_path.get().strip() or os.getcwd()
        if not url.lower().startswith(("http://", "https://")):
            messagebox.showerror("Error", "Invalid URL; must start with http:// or https://")
            return
        self._busy = True
        self.convert_btn.config(state="disabled")
        threading.Thread(target=self._run_conversion, args=(url, out_dir), daemon=True).start()

    def _run_conversion(self, url: str, out_dir: str):
        driver = None
        temp_dir = None
        ok, msg = False, ""
        try:
            self.set_status("Working\u2026")
            self.ensure_dirs(out_dir)
            temp_dir = tempfile.mkdtemp(prefix="concord_")
            self.set_progress(8)

            info = self.extract_activity_info(url)
            self.log(f"[info] detected {info['format']} format")

            data = self.fetch_source_json(info)
            self.set_progress(18)

            title, pages = self.derive_metadata(info, data)
            safe_title = self.sanitize_filename(title)

            want_pdf = self.pdf_var.get()
            want_gdocs = self.gdocs_var.get()
            want_word = self.worddoc_var.get()
            pdf_with_answers = (self.pdf_answers_var.get() == "with")

            if not (want_pdf or want_gdocs or want_word):
                raise RuntimeError("Please select at least one output (PDF, Google Docs, or Word).")

            # A PDF render is needed if the user wants the PDF itself OR the Word
            # DOCX (which is built from the PDF). The answer keys are injected only
            # when the user keeps the PDF AND chooses the "with answers" option.
            need_pdf_render = want_pdf or want_word
            inject_answers = want_pdf and pdf_with_answers

            mcq_map = collect_mcqs(data) if inject_answers else {}

            outputs = []
            pdf_path = None

            driver = self.get_driver()

            if need_pdf_render:
                self.set_progress(35)
                pdf_path = self.render_and_merge(
                    driver, info, pages, safe_title, out_dir, temp_dir,
                    mcq_map, inject_answers
                )
                if want_pdf:
                    self.log("[ok] PDF complete")
                    outputs.append(f"PDF:\n{pdf_path}")
                self.set_progress(60)

            if want_gdocs:
                self.set_status("Working\u2026")
                self.set_progress(70)
                gdocs_path, _ = self.build_docx_from_json(
                    data, info, title, out_dir, driver=driver
                )
                self.log("[ok] Google Docs version complete")
                outputs.append(f"Google Docs version:\n{gdocs_path}")
                self.set_progress(82)

            if want_word:
                self.set_status("Working\u2026")
                self.set_progress(88)
                word_path, _ = self.build_docx_highfidelity(pdf_path, title, out_dir)
                self.log("[ok] Word DOCX complete")
                outputs.append(f"Word DOCX:\n{word_path}")

            self.set_progress(100)
            msg = "Saved:\n\n" + "\n\n".join(outputs)
            ok = True

        except Exception as e:
            self.log(f"[error] {e}")
            msg = str(e)
            self.set_progress(0)
        finally:
            if driver is not None:
                try:
                    driver.quit()
                except Exception:
                    pass
            if temp_dir and os.path.isdir(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
            self.root.after(0, self._on_done, ok, msg)

    def set_progress(self, value: int):
        """Thread-safe progress bar update."""
        self.root.after(0, lambda: self.progress_bar.config(value=value))

    def _on_done(self, ok: bool, msg: str):
        self._busy = False
        self.convert_btn.config(state="normal")
        if ok:
            self.status_label.config(text="Complete!", fg=self.COLOR_SUCCESS)
            messagebox.showinfo("Success", msg)
        else:
            self.status_label.config(text="Failed", fg="#F44336")
            self.progress_bar['value'] = 0
            messagebox.showerror("Conversion failed", msg)
        # Reset for next conversion
        self.root.after(3000, self._reset_progress)

    def _reset_progress(self):
        """Reset progress bar and status to initial state."""
        if not self._busy:
            self.progress_bar['value'] = 0
            self.status_label.config(text="Ready", fg=self.COLOR_SUBTEXT)


def _enable_hidpi():
    """Make the app DPI-aware on Windows so text renders crisply instead of being
    bitmap-stretched (the usual cause of 'blurry' Tk apps on scaled displays).
    Returns the monitor scaling factor (1.0 == 100%)."""
    scaling = 1.0
    if sys.platform.startswith("win"):
        import ctypes
        # Try the modern Per-Monitor-v2 API first, then fall back to older ones.
        try:
            ctypes.windll.shcore.SetProcessDpiAwareness(2)  # PER_MONITOR_AWARE
        except Exception:
            try:
                ctypes.windll.user32.SetProcessDPIAware()    # system-DPI aware
            except Exception:
                pass
        try:
            hdc = ctypes.windll.user32.GetDC(0)
            dpi = ctypes.windll.gdi32.GetDeviceCaps(hdc, 88)  # LOGPIXELSX
            ctypes.windll.user32.ReleaseDC(0, hdc)
            if dpi:
                scaling = dpi / 96.0
        except Exception:
            scaling = 1.0
    return scaling


if __name__ == "__main__":
    _scale = _enable_hidpi()
    root = tk.Tk()
    # Tell Tk how many pixels per point so fonts are sized for this monitor.
    try:
        root.tk.call("tk", "scaling", _scale * 1.3333)
    except Exception:
        pass
    ConcordConverter(root, ui_scale=_scale)
    root.mainloop()
